import pytest
from docx import Document
from pathlib import Path
import sys
sys.path.insert(0, str(Path(__file__).parent.parent))

from docx_to_md import convert_document


def test_plain_paragraph():
    doc = Document()
    doc.add_paragraph("Hello world")
    result = convert_document(doc)
    assert result == "Hello world\n"


def test_heading():
    doc = Document()
    doc.add_heading("Title", level=1)
    doc.add_heading("Subtitle", level=2)
    result = convert_document(doc)
    assert result == "# Title\n\n## Subtitle\n"


def test_unordered_list():
    doc = Document()
    doc.add_paragraph("Item 1", style="List Bullet")
    result = convert_document(doc)
    assert result == "- Item 1\n"


def test_ordered_list():
    doc = Document()
    doc.add_paragraph("First", style="List Number")
    result = convert_document(doc)
    assert result == "1. First\n"


def test_bold_and_italic():
    doc = Document()
    p = doc.add_paragraph()
    run = p.add_run("bold text")
    run.bold = True
    run2 = p.add_run(" and ")
    run3 = p.add_run("italic text")
    run3.italic = True
    result = convert_document(doc)
    assert result == "**bold text** and _italic text_\n"


def _make_simple_table(doc):
    """Helper to create a 3x2 table with a header row, no merged cells."""
    table = doc.add_table(rows=3, cols=2)
    table.cell(0, 0).text = "A"
    table.cell(0, 1).text = "B"
    table.cell(1, 0).text = "1"
    table.cell(1, 1).text = "2"
    table.cell(2, 0).text = "3"
    table.cell(2, 1).text = "4"
    return table


def _make_complex_table_with_colspan(doc):
    """Helper to create a table with colspan."""
    from docx.oxml.ns import qn
    table = doc.add_table(rows=2, cols=3)
    table.cell(0, 0).text = "A merged"
    table.cell(0, 2).text = "C"
    table.cell(1, 0).text = "D"
    table.cell(1, 1).text = "E"
    table.cell(1, 2).text = "F"
    # Set colspan on cell (0,0) to span 2 columns
    tc = table.cell(0, 0)._tc
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        tcPr = tc.makeelement(qn('w:tcPr'), {})
        tc.insert(0, tcPr)
    gridSpan = tcPr.makeelement(qn('w:gridSpan'), {qn('w:val'): '2'})
    tcPr.append(gridSpan)
    return table


def test_table_is_simple():
    doc = Document()
    _make_simple_table(doc)
    table = doc.tables[0]
    from docx_to_md import _is_simple_table
    assert _is_simple_table(table) is True


def test_table_with_colspan_is_complex():
    doc = Document()
    _make_complex_table_with_colspan(doc)
    table = doc.tables[0]
    from docx_to_md import _is_simple_table
    assert _is_simple_table(table) is False


def test_pipe_table():
    doc = Document()
    table = doc.add_table(rows=3, cols=2)
    table.cell(0, 0).text = "Name"
    table.cell(0, 1).text = "Age"
    table.cell(1, 0).text = "Alice"
    table.cell(1, 1).text = "30"
    table.cell(2, 0).text = "Bob"
    table.cell(2, 1).text = "25"
    from docx_to_md import _convert_table_pipe
    result = _convert_table_pipe(table)
    expected = (
        "| Name  | Age |\n"
        "| ----- | --- |\n"
        "| Alice | 30  |\n"
        "| Bob   | 25  |\n"
    )
    assert result == expected


def test_html_table_with_colspan():
    from docx.oxml.ns import qn
    doc = Document()
    table = doc.add_table(rows=2, cols=3)
    # Merge first two cells in row 0
    cell = table.cell(0, 0)
    tc = cell._tc
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        tcPr = tc.makeelement(qn('w:tcPr'), {})
        tc.insert(0, tcPr)
    gridSpan = tcPr.makeelement(qn('w:gridSpan'), {qn('w:val'): '2'})
    tcPr.append(gridSpan)
    table.cell(0, 2).text = "C"
    table.cell(1, 0).text = "D"
    table.cell(1, 1).text = "E"
    table.cell(1, 2).text = "F"

    from docx_to_md import _convert_table_html
    result = _convert_table_html(table)

    assert "<table>" in result
    assert "colspan=\"2\"" in result
    assert result.strip().endswith("</table>")


def test_html_table_no_header():
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "A"
    table.cell(0, 1).text = "B"
    table.cell(1, 0).text = "C"
    table.cell(1, 1).text = "D"
    from docx_to_md import _convert_table_html
    result = _convert_table_html(table)
    assert "<table>" in result
    assert result.count("<tr>") == 2
    assert result.count("<td>") == 4


def test_mixed_content():
    doc = Document()
    doc.add_heading("Report", level=1)
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "X"
    table.cell(0, 1).text = "Y"
    table.cell(1, 0).text = "1"
    table.cell(1, 1).text = "2"
    doc.add_paragraph("Footer text")
    result = convert_document(doc)
    assert result.startswith("# Report\n")
    assert "| X | Y |" in result
    assert result.strip().endswith("Footer text")


def test_complex_table_uses_html():
    from docx.oxml.ns import qn
    doc = Document()
    doc.add_paragraph("Before")
    table = doc.add_table(rows=2, cols=2)
    cell = table.cell(0, 0)
    tc = cell._tc
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        tcPr = tc.makeelement(qn('w:tcPr'), {})
        tc.insert(0, tcPr)
    gs = tcPr.makeelement(qn('w:gridSpan'), {qn('w:val'): '2'})
    tcPr.append(gs)
    table.cell(1, 0).text = "A"
    table.cell(1, 1).text = "B"
    doc.add_paragraph("After")
    result = convert_document(doc)
    assert "Before" in result
    assert "<table>" in result
    assert "After" in result


def test_pipe_table_empty_cell():
    doc = Document()
    table = doc.add_table(rows=2, cols=3)
    table.cell(0, 0).text = "A"
    table.cell(0, 1).text = ""
    table.cell(0, 2).text = "C"
    table.cell(1, 0).text = "1"
    table.cell(1, 1).text = "2"
    table.cell(1, 2).text = "3"
    from docx_to_md import _convert_table_pipe
    result = _convert_table_pipe(table)
    assert "| A |   | C |" in result
