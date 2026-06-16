import sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


def convert_document(doc: Document) -> str:
    lines = []
    body = doc.element.body

    para_by_elem = {p._element: p for p in doc.paragraphs}
    table_by_elem = {t._tbl: t for t in doc.tables}

    for child in body:
        if child.tag == qn('w:p'):
            p = para_by_elem.get(child)
            if p is not None:
                lines.append(_convert_paragraph(p))
        elif child.tag == qn('w:tbl'):
            table = table_by_elem.get(child)
            if table is not None:
                if _is_simple_table(table):
                    lines.append(_convert_table_pipe(table))
                else:
                    lines.append(_convert_table_html(table))

    return "\n".join(lines)


def _convert_paragraph(p) -> str:
    style = p.style.name if p.style else "Normal"

    # Heading
    if style.startswith("Heading"):
        level = style.split()[-1]
        if level.isdigit():
            prefix = "#" * int(level)
        else:
            prefix = ""
        return f"{prefix} {_extract_inline_text(p)}\n"

    # Unordered list
    if style.startswith("List Bullet"):
        return f"- {_extract_inline_text(p)}\n"

    # Ordered list
    if style.startswith("List Number"):
        return f"1. {_extract_inline_text(p)}\n"

    # Normal paragraph
    return f"{_extract_inline_text(p)}\n"


def _extract_inline_text(p) -> str:
    parts = []
    for run in p.runs:
        text = run.text
        if run.bold:
            text = f"**{text}**"
        if run.italic:
            text = f"_{text}_"
        parts.append(text)
    return "".join(parts)


def _is_simple_table(table) -> bool:
    """Table is simple if it has no colspan/rowspan in any cell."""
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tc_pr = tc.find(qn('w:tcPr'))
            if tc_pr is not None:
                # colspan
                if tc_pr.find(qn('w:gridSpan')) is not None:
                    return False
                # rowspan (vMerge)
                if tc_pr.find(qn('w:vMerge')) is not None:
                    return False
    return True


def _convert_table_pipe(table) -> str:
    rows_data = []
    for row in table.rows:
        cells = []
        for cell in row.cells:
            text = _cell_text(cell)
            cells.append(text)
        rows_data.append(cells)

    if not rows_data:
        return ""

    # Column widths for alignment padding
    col_widths = []
    for col_idx in range(len(rows_data[0])):
        widths = []
        for row in rows_data:
            if col_idx < len(row):
                widths.append(len(row[col_idx]))
        col_widths.append(max(widths) if widths else 3)

    lines = []

    # Header row
    header = rows_data[0]
    line = "| " + " | ".join(h.ljust(col_widths[i]) for i, h in enumerate(header)) + " |"
    lines.append(line)

    # Separator row
    sep = "| " + " | ".join("-" * max(w, 3) for w in col_widths) + " |"
    lines.append(sep)

    # Data rows
    for row in rows_data[1:]:
        line = "| " + " | ".join(
            (row[i] if i < len(row) else "").ljust(col_widths[i])
            for i in range(len(col_widths))
        ) + " |"
        lines.append(line)

    return "\n".join(lines) + "\n"


def _cell_text(cell) -> str:
    return " ".join(p.text or "" for p in cell.paragraphs).strip()


def _get_colspan(cell) -> int:
    tc = cell._tc
    tc_pr = tc.find(qn('w:tcPr'))
    if tc_pr is None:
        return 1
    gs = tc_pr.find(qn('w:gridSpan'))
    if gs is None:
        return 1
    return int(gs.get(qn('w:val'), 1))


def _has_rowspan_restart(cell) -> bool:
    tc = cell._tc
    tc_pr = tc.find(qn('w:tcPr'))
    if tc_pr is None:
        return False
    vm = tc_pr.find(qn('w:vMerge'))
    if vm is None:
        return False
    val = vm.get(qn('w:val'))
    return val == "restart"


def _is_rowspan_continue(cell) -> bool:
    tc = cell._tc
    tc_pr = tc.find(qn('w:tcPr'))
    if tc_pr is None:
        return False
    vm = tc_pr.find(qn('w:vMerge'))
    if vm is None:
        return False
    val = vm.get(qn('w:val'))
    return val is None


def _count_rowspan(table, start_row_idx, col_idx) -> int:
    count = 1
    for row_idx in range(start_row_idx + 1, len(table.rows)):
        row = table.rows[row_idx]
        cells = row.cells
        if col_idx < len(cells) and _is_rowspan_continue(cells[col_idx]):
            count += 1
        else:
            break
    return count


def _escape_html(text: str) -> str:
    return (
        text.replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
    )


def _convert_table_html(table) -> str:
    lines = ["<table>"]
    rowspan_skip = {}

    for row_idx, row in enumerate(table.rows):
        lines.append("<tr>")
        skip_cols = set()
        for c in list(rowspan_skip.keys()):
            rowspan_skip[c] -= 1
            if rowspan_skip[c] <= 0:
                del rowspan_skip[c]
            else:
                skip_cols.add(c)
        col_idx = 0
        cell_idx = 0
        while cell_idx < len(row.cells):
            if col_idx in skip_cols:
                col_idx += 1
                continue
            cell = row.cells[cell_idx]
            colspan = _get_colspan(cell)
            text = _cell_text(cell)
            attrs = ""
            if colspan > 1:
                attrs += f' colspan="{colspan}"'
            if _has_rowspan_restart(cell):
                rs = _count_rowspan(table, row_idx, cell_idx)
                if rs > 1:
                    attrs += f' rowspan="{rs}"'
                    rowspan_skip[col_idx] = rs
            lines.append(f"  <td{attrs}>{_escape_html(text)}</td>")
            col_idx += colspan
            cell_idx += 1
        lines.append("</tr>")
    lines.append("</table>")
    return "\n".join(lines) + "\n"


def convert_docx_to_markdown(input_path: Path, output_path: Path) -> None:
    if not input_path.exists():
        raise FileNotFoundError(f"Файл не найден: {input_path}")

    if input_path.suffix.lower() != ".docx":
        raise ValueError("Поддерживается только формат .docx")

    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document(str(input_path))
    md_content = convert_document(doc)

    output_path.write_text(md_content, encoding="utf-8")


def main():
    if len(sys.argv) != 3:
        print("Usage: python docx_to_md.py input.docx output.md")
        sys.exit(1)

    input_file = Path(sys.argv[1])
    output_file = Path(sys.argv[2])

    try:
        convert_docx_to_markdown(input_file, output_file)
        print(f"Успешно сконвертировано: {output_file}")
    except Exception as e:
        print(f"Ошибка: {e}")
        sys.exit(2)


if __name__ == "__main__":
    main()
